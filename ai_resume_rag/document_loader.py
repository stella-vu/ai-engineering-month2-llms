from pathlib import Path
import fitz
from docx import Document


def load_pdf_text(file_path: str) -> str:
    try:
        text_parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""


def load_docx_text(file_path: str) -> str:
    document = Document(file_path)

    paragraphs = []
    
    # Extract text from standard paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    # Extract text from tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                # Prevent duplicate text if cells are merged
                if text and text not in paragraphs:
                    paragraphs.append(text)

    # Text Boxes / Shapes (Crucial for modern templates)
    # This searches the XML tree for any text run elements hidden inside shapes  
    for element in document.element.body.iter():
        if element.tag.endswith('t'): # 't' represents a text node in openxml
            text = element.text.strip() if element.text else ""
            if text and text not in paragraphs:
                # Basic check to avoid mixing inline text duplicates
                paragraphs.append(text)

    return "\n".join(paragraphs)


def load_text_file(file_path: str) -> str:
    path = Path(file_path)

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf_text(file_path)

    if suffix == ".docx":
        return load_docx_text(file_path)

    if suffix == ".txt":
        return path.read_text(encoding="utf-8")

    raise ValueError(f"Unsupported file type: {suffix}")

